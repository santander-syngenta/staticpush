"""
Organized package for printing Word Documents off of the Data-Lake
"""
from __future__ import print_function
import time
from pprint import pprint
import psycopg2
import numpy as np
from sqlalchemy import create_engine
import pandas as pd
import math
from docx.shared import Cm, Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
import dateutil.parser as parser
from docx.table import _Cell
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pyodbc
from win32com import client
import win32com
import threading
import re
import datetime

HOST = 'deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com'
PWD = 's1030345PASS81997!'
USER = 's1030345'
PORT = '5439'
DBNAME = 'mio'

class Doc():
	"""
	Class for printing formatted Word Documents.

	Attributes:
		Engine: Engine for connecting to Data Lake SQL Server
		Format: JSON object of formatting preferences
		Instructions: Dictionary of instructions
	"""
	def __init__(self, protocol_id, shortname, font, font_size, confidential, color):
		"""
		Initializes the connection to the Data Lake and records the formatting preferences for the Word Document
		"""
		self._engine = create_engine('postgresql://s1030345:s1030345PASS81997!@deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com:5439/mio')
		self._format = {'pid': protocol_id, 'shortname':shortname, 'font':font, 'font_size': font_size, 'confidential': confidential, 'color': color}
		self.pull_data_threaded()
		
	@property
	def engine(self):
		return self._engine

	@property
	def format(self):
		return self._format

	@property
	def instructions(self):
		return self._instructions

	@property
	def instructions_conf(self):
		return self._instructions_conf

	@property
	def assessments(self):
		return self._assessments

	@property
	def trt_df(self):
		return self._trt_df

	@property
	def total_df(self):
		return self._total_df

	@property
	def overview(self):
		return self._overview

	@property
	def status_df(self):
		return self._status_df


	def pull_data_threaded(self):
		p  = self.format['pid']
		THREAD1 = threading.Thread(target = self.pull_instructions)
		THREAD2 = threading.Thread(target = self.pull_assessments)
		THREAD3 = threading.Thread(target = self.pull_treatments)
		THREAD4 = threading.Thread(target = self.pull_overview)
		threads = [THREAD1, THREAD2, THREAD3, THREAD4]
		for x in threads:
			x.start()
		for x in threads:
			x.join()


	def pull_instructions(self):
		p = self.format['pid'].upper()
		df = pd.read_sql("SELECT instructions_instructions as instructions, instructions_confidential as design_codes, instructions_studyruleset_code as study_rules FROM mio.public.mio212_trialprotocol_instructions WHERE protocol_id = \'" +p+'\';', self.engine)
		if df.empty == True:
			d = []
		elif df.empty == False and df['instructions'][0] != None:
			df = remove_blankrows(df).reset_index()
			df = df.drop(columns=['index'])
			output = {}; c = df.columns
			for x in range(len(df.loc[0])):
				column = []
				for y in range(len(df)):
					if df.iloc[y,x] == None:
						pass
					else:
						column.append(df.iloc[y,x])
				output[c[x]] = column
			d = output
		else:
			d = []
		#Data Cleansing
		if d == []:
			self._instructions = {}
			self._instructions_conf = {}
		elif d['instructions'] == []:
			self._instructions = {}
			self._instructions_conf = {}
		else:
			instructions = d['instructions'][0]
			occurences = instructions.count('\r\n'); lines = []
			subjects = ['CROPS/SURFACES','CROPS', ' CROPS', 'TARGETS', 'OBJECTIVE','OBJECTIVES','OBJECTIVE(S)','SPECIAL PROTOCOL TASKS','DATA REQUIREMENTS/ESSENTIAL DATA','DATA REQUIEMENTS/ESSENTIAL DATA','EXPERIMENTAL DESIGN AND PLOT DIMENSIONS','EXPERIMENTAL DESIGN & PLOT DIMENSIONS','MAINTANENCE DETAILS', 'MAINTENANCE DETAILS','TREATMENT DETAILS', 'ASSESSMENT TIMING SUMMARY','ASSESSMENT DETAILS','DATA REPORTING DEADLINES','OTHER NOTES','CROP DESTRUCT','DATA REPORTING GUIDELINES','SAFETY AND STEWARDSHIP OF TEST SUBSTANCES IN THIS PROTOCOL','CONTRACT RESEARCH ORGANIZATIONS','CONTRACT RESEARCH ORGANIZATIONS (NOT UNIVERSITIES)', 'CROP DESTRUCT','UNIVERSITIES','DESIGN CODES']
			for x in range(occurences):
				index = instructions.index('\r\n')
				lines.append(instructions[0:index+2])
				instructions = instructions[index+2:]
			y = 0; indexes = []; current_para = []; output = {}; current_hdr = None
			for line in lines:
				try:
					index = line.index(':')
					sub = line[:index]
					header = True
					if header == True and sub.upper() in subjects:
						indexes.append(y)
						output[current_hdr] = current_para
						current_hdr = line
						current_para = []
					else:
						current_para.append(line)
				except:
					if line.strip('\r\n').strip().upper() in subjects:
						indexes.append(y)
						output[current_hdr] = current_para
						current_hdr = line
						current_para = []
					else:
						current_para.append(line)
				y+=1
			if output != {}:
				output.pop(None)
				out2 = {}
				for x in output.keys():
					lines = output[x]
					l_copy = lines.copy()
					lines_out = []
					for line in l_copy:
						if line == ' \r\n':
							lines.remove(line)
						else:
							if '●' in line:
								line = line.replace('●','•')
							elif '●' in line:
								line = line.replace('●','•')
							lines_out.append(line.replace('\r\n','\n'))
					out2[x] = lines_out
				confidential = d['design_codes']; unique=[]; conf_out = []
				for line in confidential:
					if line not in unique:
						unique.append(line)
						conf_out.append(line.replace('\r\n','\n'))
					else:
						pass
				
				self._instructions = out2
				self._instructions_conf = conf_out
			else:
				self._instructions = {}
				self._instructions_conf = {}


	def pull_assessments(self):
		p = self.format['pid'].upper()
		df = pd.read_sql("SELECT assessments_standardevaluationname_code as se_code, assessments_standardevaluationname_name as se_name, assessments_partrated_name as part_rated, assessments_ratingtype_code as rating_type, assessments_ratingunit_code as rating_unit, assessments_sampling as sample_size, assessments_samplingunit_code as sample_unit, assessments_collection as collection_basis, assessments_collectionunit_code as collection_unit, assessments_reporting as reporting_basis, assessments_reportingunit_code as reporting_unit, assessments_numberofsubsamples as no_subsamples FROM mio.public.mio212_trialprotocol_assessments where protocol_id =\'" + p + "\';", self.engine)
		df = remove_blankrows(df)
		df = df.applymap(str)

		##Remove Trailing Zeros
		df['no_subsamples'] = df['no_subsamples'].str.replace('.0$','')
		df['sample_size'] = df['sample_size'].str.replace('.0$','')
		df['collection_basis'] = df['collection_basis'].str.replace('.0$','')
		df['reporting_basis'] = df['reporting_basis'].str.replace('.0$','')

		df['sample_size'] = df['sample_size'].replace('nan', np.nan)
		df['collection_basis'] = df['collection_basis'].replace('nan', np.nan)
		df['reporting_basis'] = df['reporting_basis'].replace('nan', np.nan)

		df.loc[df['sample_size'].notnull(), 'new_sample_size']= df['sample_size'].astype(str) + '/' + df['sample_unit']
		df.loc[df['collection_basis'].notnull(), 'new_collect'] = df['collection_basis'].astype(str) + '/' + df['collection_unit']
		df.loc[df['reporting_basis'].notnull(),'new_report'] = df['reporting_basis'].astype(str) + '/' + df['reporting_unit']
		
		df['no_subsamples'] = df['no_subsamples'].astype(str)
		df['no_subsamples'] = df['no_subsamples'].str.replace('nan', ' ')
		df['new_sample_size'] = df['new_sample_size'].fillna(' ')
		df['new_collect'] = df['new_collect'].fillna(' ')
		df['new_report'] = df['new_report'].fillna(' ')
		df = df.drop(columns=['sample_unit','collection_unit','reporting_unit', 'sample_size', 'collection_basis', 'reporting_basis'])

		df.columns = df.columns.str.replace('no_subsamples', '# of Sub-\nSamples')
		df.columns = df.columns.str.replace('new_report', 'Reporting\nBasis')
		df.columns = df.columns.str.replace('new_collect', 'Collection\nBasis')
		df.columns = df.columns.str.replace('new_sample_size', 'Sample\nSize')
		df.columns = df.columns.str.replace('part_rated', 'Part\nRated')
		df.columns = df.columns.str.replace('se_code', 'SE Code')
		df.columns = df.columns.str.replace('se_name', 'SE Name')
		df.columns = df.columns.str.replace('rating_type', 'Rating\nType')
		df.columns = df.columns.str.replace('rating_unit', 'Rating\nUnit')

		self._assessments = df 


	def pull_treatments(self):
		p = self.format['pid'].upper()
		df = pd.read_sql("SELECT treatments_treatmentlist_treatmentnumber as no, treatments_treatmentlist_treatementseq as seq_no, treatments_treatmentlist_treatementtag as trttag, treatments_treatmentlist_formconcqty as conc_qty, treatments_treatmentlist_rate as rate, treatments_treatmentlist_otherrate as other_rate, treatments_treatmentlist_applcode as appl_code, treatments_treatmentlist_check_name as check_name, treatments_treatmentlist_productamountunit_code as unit, treatments_treatmentlist_treatmenttype_code as treatment_code, treatments_treatmentlist_treatmentname_code as trt_name, treatments_treatmentlist_treatmentname_name as trt_name2, treatments_treatmentlist_formconcunit_code as conc_unit, treatments_treatmentlist_formtype_code as Form_Type, treatments_treatmentlist_rateunit_code as rate_unit, treatments_treatmentlist_otherrateunit_code as otherrate_unit, treatments_treatmentlist_appltiming_armdescription as appl_timing, treatments_treatmentlist_applmethod_armdescription as appl_method, treatments_treatmentlist_applplacement_armdescription as appl_placement, treatments_treatmentlist_treatmenttype_code as trt_type, treatments_treatmentlist_minapplication as min_app, treatments_treatmentlist_productamounttotalqty as total_qty, treatments_treatmentlist_productamountunit_code as total_units, treatments_treatmentlist_treatmentname_name as prod_description, treatments_treatmentlist_applplacement_armcode as placement, treatments_treatmentlist_appltiming_armcode as timing, treatments_treatmentlist_applmethod_armcode as method, treatments_treatmentlist_lotcode as lotcode FROM mio.public.mio212_trialprotocol_treatments WHERE protocol_id = \'" + p + "\';", self.engine)
		df = remove_blankrows(df)

		#Data Cleansing
		if df.empty == True:
			self._trt_df = df; self._total_df = {}
		else:
			df = remove_blankrows(df)
			df.columns = df.columns.str.replace('trt_name','Treatment Name')
			df2 = df[['Treatment Name','total_qty','total_units','Treatment Name2']]
			df['no'] = df['no'].fillna(0)
			df['seq_no'] = df['seq_no'].fillna(0)
			df['trttag'] = df['trttag'].fillna(0)
			df = df.fillna(' ')
			df = df.applymap(str)
			for x in range(len(df)):
				df.iloc[x,0] = str(int(float(df.iloc[x,0])))

			"""Rename & Reorder columns in Treatment Table"""
			df.columns = df.columns.str.replace('no','No.')
			df = df.drop(columns=['trt_type'])
			df.columns = df.columns.str.replace('appl_code','Code')
			df.columns = df.columns.str.replace('form_type', 'Type')
			df.columns = df.columns.str.replace('min_app', 'Min #\nAppl')
			df.columns = df.columns.str.replace('other_rate','Other\nRate')
			df.columns = df.columns.str.replace('otherrate_unit', 'Other Rate\nUnit')
			df.columns = df.columns.str.replace('rate', 'Rate')
			df.columns = df.columns.str.replace('Rate_unit','Rate\nUnit')
			df.columns = df.columns.str.replace('conc_unit', 'Form.\nUnit')
			df.columns = df.columns.str.replace('conc_qty', 'Form.')
			df.columns = df.columns.str.replace('lotcode', 'Lot\nCode')
			df.columns = df.columns.str.replace('method','Method')
			df.columns = df.columns.str.replace('placement', 'Place\nment')
			df.columns = df.columns.str.replace('timing', 'Timing')
			df.columns = df.columns.str.replace('prod_description', 'Description')
			df = df[['No.', 'trttag', 'seq_No.', 'Treatment Name', 'Description','Form.', 'Form.\nUnit', 'Type', 'Rate','Rate\nUnit','Other\nRate','Other Rate\nUnit','Place\nment','Timing','Method','Min #\nAppl','Code','Lot\nCode']]
			remove = []
			for column in df.columns:
				blank = True
				for value in df[column]:
					if value.strip():
						blank = False
						break
				if blank == True:
					remove.append(column)
			df = df.drop(columns=remove)
			df2 = df2.dropna()
			unique = df2['Treatment Name'].unique(); d = {}
			for x in unique:
				d[x] = [0,None,None]
			for x in df2.iterrows():
				d[x[1][0]][0] += round(float(x[1][1]),2)
				d[x[1][0]][1] = x[1][2]
				d[x[1][0]][2] = x[1][3]
			if 'No.' in df.columns:
				df['No.'] = df['No.'].astype(int)
			df = df.sort_values(by=['No.', 'trttag', 'seq_No.'])
			df = df.drop(columns=['seq_No.','trttag', 'Description'])
			df = df.dropna(axis=1,how='all')
			"""Removing Trailing zeros without rounding"""
			if 'Rate' in df.columns:
				df['Rate'] = df['Rate'].str.replace('.0$','')
			if 'Form.' in df.columns:
				df['Form.'] = df['Form.'].str.replace('.0$','')
			if 'Min #\nAppl' in df.columns:
				df['Min #\nAppl'] = df['Min #\nAppl'].str.replace('.0$','')

			self._trt_df = df; self._total_df = d 


	def pull_overview(self):
		p = self.format['pid'].upper()
		overview_df = pd.read_sql("SELECT overview_setup_discipline_name as discipline, overview_setup_project as project, overview_setup_stage_name as stage, overview_setup_country_iso3 as country, overview_setup_origin_code as usage_code, overview_setup_origin_name as usage, overview_users_ownedby_firstname as owner_first, overview_users_ownedby_lastname as owner_last, overview_setup_interimdataneededby as interimdataneededby, overview_setup_finaldataneededby as finaldataneededby, overview_setup_neededby as protocolneededby, overview_objectives_objectivenormal as objective, overview_objectives_objective as overview2, overview_crops_crop_name as crop, overview_users_reviewers_firstname as reviewer_first, overview_users_contributors_firstname as contributor_first, overview_users_contributors_lastname as contributor_last, overview_users_reviewers_lastname as reviewer_last, overview_setup_trialstartyear as start_year, shortname, overview_targets_target_armcode as target_code, overview_targets_target_armdescription as target, overview_crops_crop_code as crop_code, title FROM mio.public.mio212_trialprotocol_fact WHERE protocol_id =  \'" + p + "\';", self.engine)
		status_df = pd.read_sql("SELECT top 1 status, revision, lastupdateddate FROM mio.public.mio212_trialprotocol_fact WHERE protocol_id =  \'" + p + "\';", self.engine)

		if overview_df.empty == False:
			df = overview_df
			df['owner'] = df['owner_first'] + ' '+ df['owner_last']
			df['reviewer'] = df['reviewer_first'] + ' '+ df['reviewer_last']
			df['contributor'] = df['contributor_first'] + ' ' + df['contributor_last']
			df = df.drop(columns=['owner_first','owner_last', 'reviewer_first', 'reviewer_last', 'contributor_last', 'contributor_first'])
			output = {}; c = df.columns
			for x in range(len(df.loc[0])):
				column = []
				for y in range(len(df)):
					if (type(df.iloc[y,x]) == float or type(df.iloc[y,x]) == np.float64) and math.isnan(df.iloc[y,x]) == True:
						pass
					elif df.iloc[y,x] == None:
						pass
					else:
						column.append(df.iloc[y,x])
				output[c[x]] = column
			for x in range(len(output['target'])):
				t = output['target'][x] + ' (' + output['target_code'][x] +')'
				output['target'][x] = t
			for x in range(len(output['crop'])):
				c = output['crop'][x] + ' (' + output['crop_code'][x] + ')'
				output['crop'][x] = c
			output.pop('crop_code')
			output.pop('target_code')
			output['title'] = df['title'][0]
		else:
			output = []

		self._overview = output
		self._status_df = status_df.fillna(' ')


	def add_cover(self, document):
		##Declare Local Variables
		protocol_id = self.format['pid']
		d = self.overview
		confidential = self.format['confidential']
		font_type = self.format['font']
		font_size = self.format['font_size']
		status_df = self.status_df

		#Add to Word
		document.styles['Normal'].font.name = font_type
		shortname = document.add_paragraph()
		short_run = shortname.add_run("\r\n\r\n" + d['shortname'][0])
		short_run.font.bold = True
		short_run.font.size = Pt(font_size + 2)
		shortname.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		interim, final  = parser.parse(d['interimdataneededby'][0]).strftime("%a, %B %d %Y"), parser.parse(d['finaldataneededby'][0]).strftime("%a, %B %d %Y")
		cover_body = document.add_paragraph()
		title = cover_body.add_run(d['title'])
		title.font.size = Pt(font_size)
		title.font.bold = True

		info = cover_body.add_run('\r\nProtocol Owner: ' + d['owner'][0] + '\r\nContributors: ' + ', '.join(d['contributor']) + '\r\nReviewers: ' + ', '.join(d['reviewer']))
		cover_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		table1 = document.add_table(rows=3,cols=2)
		table1.style.name='Table Grid'
		table1.alignment = WD_TABLE_ALIGNMENT.CENTER
		header_row = table1.rows[0].cells
		header_row[0].text = "Interim Data Needed By:"
		header_row[1].text = "Final Data Needed By:"
		header_p = header_row[0].paragraphs[0]
		header_p2 = header_row[1].paragraphs[0]
		header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
		header_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
		header_p2.runs[0].font.bold = True
		header_p.runs[0].font.bold = True
		table1.rows[1].cells[0].text = interim
		table1.rows[1].cells[1].text = final
		table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		table1.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		try:
			status = status_df.iloc[0,0]
		except:
			status = 'N/A'
		try:
			revision = status_df.iloc[0,1]
		except:
			revision = 'N/A'
		try:
			if status_df.iloc[0,2] != ' ':
				last_update = status_df.iloc[0,2].replace(re.findall(r'.\d{7}$', status_df.iloc[0,2])[0], '') ###Removes additional decimal values that are unneccesary
				last_update_date = datetime.datetime.strptime(last_update, "%Y-%m-%dT%H:%M:%S").strftime("%b %w, %Y; %I:%M%p")
			else:
				last_update_date = 'N/A'
		except Exception as e:
			last_update_date = 'N/A'
		table2 = document.add_table(rows=2,cols=3)
		table2.style.name='Table Grid'
		table2.alignment = WD_TABLE_ALIGNMENT.CENTER
		header_row = table2.rows[0].cells
		header_row[0].text = "Protocol Status:"
		header_row[1].text = "Revision No.:"
		header_row[2].text = "Last Updated:"
		header_p = header_row[0].paragraphs[0]
		header_p2 = header_row[1].paragraphs[0]
		header_p3 = header_row[2].paragraphs[0]
		header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		header_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
		header_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
		header_p2.runs[0].font.bold = True
		header_p3.runs[0].font.bold = True
		header_p.runs[0].font.bold = True
		table2.rows[1].cells[0].text = status
		table2.rows[1].cells[1].text = revision
		table2.rows[1].cells[2].text = last_update_date
		table2.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		table2.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		table2.rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		document.add_page_break()
		header = document.sections[0].header.paragraphs[0]
		run = header.add_run()
		run.add_picture('.\\logo.jpg', width = Inches(1))
		footer = document.sections[0].footer.paragraphs[0]
		footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		run2 = footer.add_run()
		if confidential == True:
			run2.text = d['shortname'][0]
			run3 = footer.add_run()
			run3.text = "  For Internal Distribution (Confidential)"
			run3.font.color.rgb = RGBColor(255,0,0)
		else:
			run2.text = d['shortname'][0]
			run3 = footer.add_run()
			run3.text = "  Suitable for External Distribution (Confidential)"
			run3.font.color.rgb = RGBColor(255,0,0)
		header.style = document.styles['Header']
		footer.style = document.styles['Footer']
		run.alignment = WD_ALIGN_PARAGRAPH.LEFT
		add_page_number(document.sections[0].footer.paragraphs[0])


	def add_instructions(self, document):
		instructions = self.instructions
		protocol_id = self.format['pid']
		font_size = self.format['font_size']
		if self.format['confidential'] == False:
			confidential_inst = self.instructions_conf
		else:
			confidential_inst = None

		d = self.instructions
		for header in d.keys():
			content = document.add_paragraph()
			try:
				index = header.index(':')
				f = header[:index]
				s = header[index:]
				hdr = content.add_run(f)
				hdr2 = content.add_run(s.replace('\r\n','\n'))
			except:
				hdr = content.add_run(header.strip('\r\n') + ':\n')
			content.paragraph_format.keep_together = True; y = 1;
			for x in d[header]:
				if y != len(d[header]):
					try:
						index = x[:50].index(':')
						bold = x[:index]
						sub = x[index:]
						bold_run = content.add_run(bold)
						sub_run = content.add_run(sub)
						bold_run.font.bold = True
						sub_run.font.bold = False
						bold_run.font.size = Pt(font_size)
						sub_run.font.size = Pt(font_size)
					except:
						content_run = content.add_run(x)
						content_run.font.size = Pt(font_size)
						content_run.font.bold = False
				else:
					try:
						index = x[:50].index(':')
						bold = x[:index]
						sub = x[index:]
						bold_run = content.add_run(bold)
						sub_run = content.add_run(sub.strip('\r\n'))
						bold_run.font.bold = True
						sub_run.font.bold = False
						bold_run.font.size = Pt(font_size)
						sub_run.font.size = Pt(font_size)
					except:
						content_run = content.add_run(x.strip('\r\n'))
						content_run.font.size = Pt(font_size)
						content_run.font.bold = False
				y += 1
			content.paragraph_format.left_indent = Inches(0)
			content.paragraph_format.space_after = Pt(6)
			content.paragraph_format.line_spacing = 1
			hdr.font.size = Pt(font_size + 2)
			hdr.font.underline = True
			hdr.font.bold = True
		section = document.sections[-1]
		section.left_margin = Inches(0.35)
		section.right_margin = Inches(0.35)
		section.top_margin = Inches(0.5)
		section.bottom_margin = Inches(0.5)
		section.header_distance = Inches(0.35)
		section.footer_distance = Inches(0.35)
		document.add_page_break()
		if confidential_inst != None:
			conf_p = document.add_paragraph()
			conf_head = conf_p.add_run('The following information is Confidential and intendeded for internal use only. Do not share this information with external cooperator.\r\n')
			conf_head.font.bold = True
			conf_head.font.size = Pt(font_size)
			conf_head.font.color.rgb = RGBColor(255,0,0)
			for x in confidential_inst:
				r = conf_p.add_run(x)
				r.font.bold = False
				r.font.size = Pt(font_size)
			document.add_page_break()
		else:
			pass


	def add_assessments(self, document):
		protocol_id = self.format['pid']
		df = self.assessments
		font_type = self.format['font']
		font_size = self.format['font_size']
		color = self.format['color']

		title = document.add_paragraph()
		title_run = title.add_run('Assessments')
		title_run.font.bold = True
		title_run.font.size = Pt(font_size + 2)
		title_run.font.underline = True
		title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		cols = df.columns
		table = document.add_table(rows=1, cols=len(cols))
		table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
		if font_type == 'Arial':
			font_type2 = 'Arial Narrow'
		else:
			font_type2 = font_type
		style = document.styles['table']; font = style.font; font.name = font_type2; font.size = Pt(font_size); font.bold = False
		style2 = document.styles['table_header']; font = style2.font; font.name = font_type; font.size = Pt(font_size); font.bold = True; font.color.rgb = RGBColor(255,255,255)
		hdr_cells = table.rows[0].cells; y = 0
		for x in cols:
			hdr_cells[y].text = x
			p = hdr_cells[y].paragraphs[0]
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p.style = document.styles['table_header']
			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
			hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
			y+=1
		y = 0; index = 1;
		for x in df.iterrows():
			border_row = False
			row = x[1]
			row_cells = table.add_row().cells
			for i in range(len(cols)):
				row_cells[i].text = str(row[i])
				p = row_cells[i].paragraphs[0]
				p.style = document.styles['table']
				set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if y % 2 == 1:
					if color == True:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
					else:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
					row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
				if index == len(df):
					set_cell_border(row_cells[i], bottom={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if i == len(cols)-1:
					set_cell_border(row_cells[i], end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			y += 1
			index += 1
		for column in table.columns:
			for cell in column.cells:
			    tc = cell._tc
			    tcPr = tc.get_or_add_tcPr()
			    tcW = tcPr.get_or_add_tcW()
			    tcW.type = 'auto'


	def add_totals(self, document):
		"""Removes confidential trt_name info and makes table"""
		d = self.total_df
		font = self.format['font']
		font_size = self.format['font_size']
		color = self.format['color']
		confidential = self.format['confidential']

		title = document.add_paragraph()
		title_run = title.add_run('\r\nProduct Totals')
		title_run.font.bold = True
		title_run.font.size = Pt(font_size + 2)
		title_run.font.underline = True
		title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		if confidential == False:
			data = {}
			for x in d.keys():
				data[x] = d[x][0:2]
			table = document.add_table(rows=1,cols=3)
			table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
			hdr_cells = table.rows[0].cells; headers = ['Treatment Name','Product\nTotal','Unit']
			for x in range(3):
				hdr_cells[x].text = headers[x]
				p = hdr_cells[x].paragraphs[0]
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER
				p.style = document.styles['table_header']
				shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
				hdr_cells[x]._tc.get_or_add_tcPr().append(shading_elm_1)
				set_cell_border(hdr_cells[x], end={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"},start={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"})
			y = 0
			for x in data.keys():
				row_cells = table.add_row().cells
				row_cells[0].text = x
				row_cells[1].text = str(round(data[x][0],2))
				row_cells[2].text = data[x][1]
				for i in range(3):
					p = row_cells[i].paragraphs[0]
					p.style = document.styles['table']
					set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
					if y % 2 == 1:
						if color == True:
							shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
						else:
							shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
						row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
				y+=1
			for column in table.columns:
				for cell in column.cells:
				    tc = cell._tc
				    tcPr = tc.get_or_add_tcPr()
				    tcW = tcPr.get_or_add_tcW()
				    tcW.type = 'auto'
		else:
			data ={}
			for x in d.keys():
				data[x] = d[x]
			table = document.add_table(rows=1,cols=4)
			table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
			hdr_cells = table.rows[0].cells; headers = ['Treatment Name','Product Description','Product\nTotal','Unit']
			for x in range(4):
				hdr_cells[x].text = headers[x]
				p = hdr_cells[x].paragraphs[0]
				p.alignment = WD_ALIGN_PARAGRAPH.CENTER
				p.style = document.styles['table_header']
				shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
				hdr_cells[x]._tc.get_or_add_tcPr().append(shading_elm_1)
				set_cell_border(hdr_cells[x], end={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"},start={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"})
			y = 0
			for x in data.keys():
				row_cells = table.add_row().cells
				row_cells[0].text = x
				row_cells[1].text = data[x][2]
				row_cells[2].text = str(round(data[x][0],2))
				row_cells[3].text = data[x][1]
				for i in range(4):
					p = row_cells[i].paragraphs[0]
					p.style = document.styles['table']
					set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
					if y % 2 == 1:
						if color == True:
							shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
						else:
							shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
						row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
				y+=1
			for column in table.columns:
				for cell in column.cells:
				    tc = cell._tc
				    tcPr = tc.get_or_add_tcPr()
				    tcW = tcPr.get_or_add_tcW()
				    tcW.type = 'auto'


	def add_treatments(self, document):
		protocol_id = self.format['pid']
		df = self.trt_df
		table_font = self.format['font']
		font_size = self.format['font_size']
		color = self.format['color']
		confidential = self.format['confidential']

		title = document.add_paragraph()
		title_run = title.add_run('Treatments')
		title_run.font.bold = True
		title_run.font.size = Pt(font_size + 2)
		title_run.font.underline = True
		title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
		cols = df.columns
		table = document.add_table(rows=1,cols=len(cols))
		table.allow_autofit = True
		table.alignment = WD_TABLE_ALIGNMENT.CENTER
		hdr_cells = table.rows[0].cells; y = 0
		for x in cols:
			hdr_cells[y].text = x
			p = hdr_cells[y].paragraphs[0]
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p.style = document.styles['table_header']
			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
			hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
			y+=1
		y = 0; current = 0; shaded = True
		for x in df.iterrows():
			border_row = False
			row = x[1]
			row_cells = table.add_row().cells
			last = False
			if y == len(df)-1:
				last = True
			for i in range(len(cols)): 
				row_cells[i].text = str(row[i])
				p = row_cells[i].paragraphs[0]
				p.style = document.styles['table']
				set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},)
				if i == 0:
					if int(row[i]) > current:
						current = int(row[i])
						border_row = True
						if shaded == False:
							shaded = True
						elif shaded == True:
							shaded = False
				elif i == len(cols)-1:
					set_cell_border(row_cells[i],end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if border_row == True:
					set_cell_border(row_cells[i], top={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if shaded == True:
					if color == True:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
					else:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
					row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
				if last == True:
					set_cell_border(row_cells[i],bottom={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			y += 1
		for column in table.columns:
			for cell in column.cells:
			    tc = cell._tc
			    tcPr = tc.get_or_add_tcPr()
			    tcW = tcPr.get_or_add_tcW()
			    tcW.type = 'auto'
		table.style.name='Table Grid'
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width 
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = new_width
		section.page_height = new_height
		section.left_margin = Inches(0.5)
		section.right_margin = Inches(0.5)
		section.top_margin = Inches(0.5)
		section.bottom_margin = Inches(0.5)
		section.header_distance = Inches(0.35)
		section.footer_distance = Inches(0.35)


	def printer(self):
		"""
		Prints to a Formatted Word Document & saves to disk.
		"""
		document = Document()
		##Check if overview or Protocol status data is blank. If not then it add the cover sheet.
		if self.overview != [] and self.status_df.empty == False:
			self.add_cover(document)
		#Checks if instructions are blank
		if self.instructions != {}:
			self.add_instructions(document)
		###Start new landscape section and adds in the assessments, product totals, and treatments if they are not blank. 
		document.add_section(WD_SECTION.NEW_PAGE)
		self.add_styles(document)
		if self.assessments.empty == False:
			self.add_assessments(document)
		if self.total_df != {}:
			self.add_totals(document)
		if self.trt_df.empty == False:
			document.add_page_break()
			self.add_treatments(document)

		document.save('.\\Q\\' + self.format['shortname'] + '.docx')
		#document.save('C:\\inetpub\\wwwroot\\nematool\\static\\docs\\' + self.format['shortname'] + '.docx')


	def add_styles(self, document):
		font_type = self.format['font']
		font_size = self.format['font_size']
		if font_type == 'Arial':
			font_type2 = 'Arial Narrow'
		elif font_type == 'Cambria' or 'Times New Roman':
			font_type2 = 'Cambria Math'
		else:
			font_type2 = font_type
		style = document.styles.add_style('table', WD_STYLE_TYPE.PARAGRAPH); font = style.font; font.name = font_type2; font.size = Pt(font_size-2); font.bold = False
		style2 = document.styles.add_style('table_header', WD_STYLE_TYPE.PARAGRAPH); font = style2.font; font.name = font_type; font.size = Pt(font_size); font.bold = True; font.color.rgb = RGBColor(255,255,255)

def remove_blankrows(df):
	rows = []
	for x in range(len(df)):
		blank = True
		for y in range(len(df.iloc[x])):
			if (type(df.iloc[x,y]) == float and math.isnan(df.iloc[x,y]) == True) or (type(df.iloc[x,y]) == np.float64 and math.isnan(df.iloc[x,y]) == True) or df.iloc[x,y] == None:
				pass
			else:
				blank = False
		if blank == True:
			rows.append(x)
	df = df.drop(index=rows)
	return df

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):

    page_run = paragraph.add_run('    ')
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)