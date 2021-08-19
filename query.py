import psycopg2
import numpy as np
from sqlalchemy import create_engine
import pandas as pd
import math
from docx.shared import Cm, Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
import dateutil.parser as parser
from docx.table import _Cell
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


host = 'deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com'
pwd = 's1030345PASS81997!'
host = 'deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com'
user = 's1030345'
port = '5439'
dbname = 'mio'

def simple_connect():
	connection = psycopg2.connect(dbname=dbname,host=host,password=pwd,port=port,user=user)
	cursor = connection.cursor()
	return connection, cursor


def pd_connect():
	engine = create_engine('postgresql://s1030345:s1030345PASS81997!@deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com:5439/mio')
	return engine


def pd_overview(protocol_id):
	"""Returns a DataFrame with the relevant fields from the overview table. Runs slower but retains table's column headers."""
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT overview_setup_discipline_name as discipline, overview_setup_project as project, overview_setup_stage_name as stage, overview_setup_country_iso3 as country, overview_setup_origin_code as usage_code, overview_setup_origin_name as usage, overview_users_ownedby_firstname as owner_first, overview_users_ownedby_lastname as owner_last, overview_setup_interimdataneededby as interimdataneededby, overview_setup_finaldataneededby as finaldataneededby, overview_setup_neededby as protocolneededby, overview_objectives_objectivenormal as objective, overview_objectives_objective as overview2, overview_crops_crop_name as crop, overview_users_reviewers_firstname as reviewer_first, overview_users_contributors_firstname as contributor_first, overview_users_contributors_lastname as contributor_last, overview_users_reviewers_lastname as reviewer_last, overview_setup_trialstartyear as start_year, shortname, overview_targets_target_armcode as target_code, overview_targets_target_armdescription as target, overview_crops_crop_code as crop_code FROM mio.public.mio212_trialprotocol_overview WHERE protocol_id =  \'" + p + "\';", engine)
	df2 = pd.read_sql("SELECT balancing_proposedprotocolname FROM mio.public.mio189_biobalance_fact where derivedprotocol_id= \'" + p + "\';", engine)
	return df, df2


def get_applications(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT treatments_trialdesign_numberofreplicates as no_replicates, treatments_trialdesign_treatedplotaream2 as treated_area,treatments_applicationlist_code as app_code, treatments_applicationlist_volumemin as min_vol, treatments_applicationlist_volumemax as max_vol, treatments_applicationlist_requiredmixsize as mix_size, treatments_applicationlist_percentageoverage as perc_overage, treatments_applicationlist_adjustedmixsize as adjusted_size, treatments_applicationlist_totalarea as total_area, treatments_applicationlist_method_armdescription as app_method, treatments_applicationlist_timing_armdescription as timing, treatments_applicationlist_placement_armdescription as placement, treatments_applicationlist_volumeunit_code as volume_unit, treatments_applicationlist_mixsizeunit_code as mix_unit from mio.public.mio212_trialprotocol_treatments WHERE protocol_id = \'" + p + "\';", engine)
	rows_to_drop = []
	for i, x in df.iterrows():
		if x['app_code'] == None:
			rows_to_drop.append(i)
	df = df.drop(index=rows_to_drop)
	df = df.applymap(str)
	df['mix_size'] = df['mix_size'] + ' ' + df['mix_unit']
	df['min_vol'] = df['min_vol'] + ' ' + df['volume_unit']
	df['max_vol'] = df['max_vol'] + ' ' + df['volume_unit']
	df = df.drop(columns=['mix_unit', 'volume_unit'])
	return df

def get_assessments(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT assessments_standardevaluationname_code as se_code, assessments_standardevaluationname_name as se_name, assessments_partrated_name as part_rated, assessments_ratingtype_code as rating_type, assessments_ratingunit_code as rating_unit, assessments_sampling as sample_size, assessments_samplingunit_code as sample_unit, assessments_collection as collection_basis, assessments_collectionunit_code as collection_unit, assessments_reporting as reporting_basis, assessments_reportingunit_code as reporting_unit, assessments_numberofsubsamples as no_subsamples FROM mio.public.mio212_trialprotocol_assessments where protocol_id =\'" + p + "\';", engine)
	df = remove_blankrows(df)
	df['no_subsamples'] = df['no_subsamples'].fillna(0).map(int)
	df['sample_size'] = df['sample_size'].fillna(0).map(int)
	df['collection_basis'] = df['collection_basis'].fillna(0).map(int)
	df['reporting_basis'] = df['reporting_basis'].fillna(0).map(int)
	df = df.applymap(str)
	df['sample_size'] = df['sample_size'] + '/' + df['sample_unit']
	df['collection_basis'] = df['collection_basis'] + '/' + df['collection_unit']
	df['reporting_basis'] = df['reporting_basis'] + '/' + df['reporting_unit']
	df = df.drop(columns=['sample_unit','collection_unit','reporting_unit'])
	df.columns = df.columns.str.replace('no_subsamples', '# of\nSubsamples')
	df.columns = df.columns.str.replace('reporting_basis', 'Reporting\nBasis')
	df.columns = df.columns.str.replace('collection_basis', 'Collection\nBasis')
	df.columns = df.columns.str.replace('sample_size', 'Sample\nSize')
	df.columns = df.columns.str.replace('part_rated', 'Part\nRated')
	df.columns = df.columns.str.replace('se_code', 'SE Code')
	df.columns = df.columns.str.replace('se_name', 'SE Name')
	df.columns = df.columns.str.replace('rating_type', 'Rating\nType')
	df.columns = df.columns.str.replace('rating_unit', 'Rating\nUnit')
	return df


def get_treatments(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT treatments_treatmentlist_treatmentnumber as no, treatments_treatmentlist_formconcqty as conc_qty, treatments_treatmentlist_rate as rate, treatments_treatmentlist_otherrate as other_rate, treatments_treatmentlist_applcode as appl_code, treatments_treatmentlist_check_name as check_name, treatments_treatmentlist_productamountunit_code as unit, treatments_treatmentlist_treatmenttype_code as treatment_code, treatments_treatmentlist_treatmentname_code as trt_name, treatments_treatmentlist_treatmentname_name as trt_name2, treatments_treatmentlist_formconcunit_code as conc_unit, treatments_treatmentlist_formtype_code as Form_Type, treatments_treatmentlist_rateunit_code as rate_unit, treatments_treatmentlist_otherrateunit_code as otherrate_unit, treatments_treatmentlist_appltiming_armdescription as appl_timing, treatments_treatmentlist_applmethod_armdescription as appl_method, treatments_treatmentlist_applplacement_armdescription as appl_placement, treatments_treatmentlist_treatmenttype_code as trt_type, treatments_treatmentlist_minapplication as min_app FROM mio.public.mio212_trialprotocol_treatments WHERE protocol_id = \'" + p + "\';", engine)
	df = remove_blankrows(df)
	return df


def get_instructions(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT instructions_instructions as instructions, instructions_confidential as design_codes, instructions_studyruleset_code as strudy_rules FROM mio.public.mio212_trialprotocol_instructions WHERE protocol_id = \'" +p+'\';', engine)
	df = remove_blankrows(df)
	output = {}; c = df.columns
	for x in range(len(df.loc[0])):
		column = []
		for y in range(len(df)):
			if df.iloc[y,x] == None:
				pass
			else:
				column.append(df.iloc[y,x])
		output[c[x]] = column
	return output
	

def clean_instructions(protocol_id):
	d = get_instructions(protocol_id)
	instructions = d['instructions'][0]
	occurences = instructions.count('\r\n'); lines = []
	for x in range(occurences):
		index = instructions.index('\r\n')
		lines.append(instructions[0:index])
		instructions = instructions[index+2:]
	y = 0; indexes = []; current_para = []; output = {}; current_hdr = None
	for line in lines:
		try:
			index = line.index(':')
			sub = line[:index]
			header = True
			for x in sub:
				if x.isalpha() == True and x.isupper() == True:
					header = True 
				elif x.isalpha() == True and x.isupper() == False:
					header = False
					break 
				else:
					pass
			if header == True:
				indexes.append(y)
				output[current_hdr] = current_para
				current_hdr = line
				current_para = []
			else:
				current_para.append(line)
		except:
			current_para.append(line)
		y+=1
	for x in output.keys():
		lines = output[x]
		for line in lines:
			if line == ' ':
				lines.remove(line)
			else:
				line = line.strip('\r\n')
	return lines, indexes, output


def get_overview(protocol_id):
	"""Returns a DataFrame with the relevant fields from the overview table. Runs faster but no column headers."""
	connection, cursor = simple_connect()
	p = protocol_id.upper()
	cursor.execute("SELECT overview_setup_discipline_name as discipline, overview_setup_project as project, overview_setup_stage_name as stage, overview_setup_country_name as country, overview_setup_origin_code as usage_code, overview_setup_origin_name as usage, overview_users_ownedby_firstname as owner_first, overview_users_ownedby_lastname as owner_last, overview_setup_interimdataneededby as interimdataneededby, overview_setup_finaldataneededby as finaldataneededby, overview_setup_neededby as protocolneededby, overview_objectives_objectivenormal as objective, overview_crops_crop_name as crop, overview_users_reviewers_firstname as reviewer_first, overview_users_reviewers_lastname FROM mio.public.mio212_trialprotocol_overview WHERE protocol_id = \'" + p + "\';")
	data = np.array(cursor.fetchall())
	cursor.close(); connection.close()
	return data


def condense_pdoverview(df, df2):
	df['owner'] = df['owner_first'] + ' '+ df['owner_last']
	df['reviewer'] = df['reviewer_first'] + ' '+ df['reviewer_last']
	df['contributor'] = df['contributor_first'] + ' ' + df['contributor_last']
	df = df.drop(columns=['owner_first','owner_last', 'reviewer_first', 'reviewer_last', 'contributor_last', 'contributor_first'])
	output = {}; c = df.columns
	for x in range(len(df.loc[0])):
		column = []
		for y in range(len(df)):
			if (type(df.iloc[y,x]) == float or type(df.iloc[y,x]) == np.float64) and math.isnan(df.iloc[y,x]) == True:
				pass
			elif df.iloc[y,x] == None:
				pass
			else:
				column.append(df.iloc[y,x])
		output[c[x]] = column
	for x in range(len(output['target'])):
		t = output['target'][x] + ' (' + output['target_code'][x] +')'
		output['target'][x] = t
	for x in range(len(output['crop'])):
		c = output['crop'][x] + ' (' + output['crop_code'][x] + ')'
		output['crop'][x] = c
	output.pop('crop_code')
	output.pop('target_code')
	output['title'] = df2.iloc[0,0]
	return output


def remove_blankrows(df):
	rows = []
	for x in range(len(df)):
		blank = True
		for y in range(len(df.iloc[x])):
			if (type(df.iloc[x,y]) == float and math.isnan(df.iloc[x,y]) == True) or (type(df.iloc[x,y]) == np.float64 and math.isnan(df.iloc[x,y]) == True) or df.iloc[x,y] == None:
				pass
			else:
				blank = False
		if blank == True:
			rows.append(x)
	df = df.drop(index=rows)
	return df


def clean_trtdf(protocol_id):
	df = get_treatments(protocol_id)
	df = remove_blankrows(df)
	df = df.fillna(' ')
	df = df.applymap(str)
	df['Formulation'] = df['conc_qty'] + '   ' + df['conc_unit']
	df['Rate'] = df['rate'] + '   ' + df['rate_unit']
	df['Other Rate'] = df['other_rate'] + '   ' + df['otherrate_unit']
	df = df.drop(columns=['conc_qty','conc_unit','rate','rate_unit','otherrate_unit','other_rate'])
	df = df[['no','trt_type','trt_name','Formulation','form_type','Rate','Other Rate','min_app','appl_code']]
	for x in range(len(df)):
		df.iloc[x,0] = str(int(float(df.iloc[x,0])))
	df.columns = df.columns.str.replace('no','No.')
	df.columns = df.columns.str.replace('trt_type','Type')
	df.columns = df.columns.str.replace('trt_name','Treatment Name')
	df.columns = df.columns.str.replace('appl_code','Code')
	df.columns = df.columns.str.replace('form_type', 'Type')
	df.columns = df.columns.str.replace('min_app', 'Min #\nAppl')
	remove = []
	for column in df.columns:
		blank = True
		for value in df[column]:
			if value.strip():
				blank = False
				break
		if blank == True:
			remove.append(column)
	df = df.drop(columns=remove)
	return df


def print(protocol_id):
	print_doc(protocol_id, 'Cambria', 12, True)


def print_doc(protocol_id, font, font_size, confidential):
	df = clean_trtdf(protocol_id)
	df2, df3 = pd_overview(protocol_id)
	overview = condense_pdoverview(df2,df3)
	assessments = get_assessments(protocol_id)
	d = Document()
	add_cover(d, protocol_id, overview, confidential, font, font_size)
	add_assessments(d, protocol_id, assessments, font, font_size)
	add_instructions(d, protocol_id)
	add_trttable(d, protocol_id, df, font, font_size)


def add_cover(document, protocol_id, d, confidential, font_type, font_size):
	document.styles['Normal'].font.name = font_type
	shortname = document.add_paragraph()
	short_run = shortname.add_run("\r\n\r\n" + d['shortname'][0])
	short_run.font.bold = True
	short_run.font.size = Pt(font_size + 2)
	shortname.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	interim, final  = parser.parse(d['interimdataneededby'][0]).strftime("%A, %B %d %Y"), parser.parse(d['finaldataneededby'][0]).strftime("%A, %B %d %Y")
	cover_body = document.add_paragraph()
	title = cover_body.add_run(d['title'])
	title.font.size = Pt(font_size)
	title.font.bold = True

	info = cover_body.add_run('\r\nProtocol Owner: ' + d['owner'][0] + '\r\nContributors: ' + ', '.join(d['contributor']) + '\r\nReviewers: ' + ', '.join(d['reviewer']))
	cover_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	table1 = document.add_table(rows=2,cols=2)
	table1.style.name='Table Grid'
	table1.alignment = WD_TABLE_ALIGNMENT.CENTER
	header_row = table1.rows[0].cells
	header_row[0].text = "Interim Data Needed By:"
	header_row[1].text = "Final Data Needed By:"
	header_p = header_row[0].paragraphs[0]
	header_p2 = header_row[1].paragraphs[0]
	header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	header_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
	header_p2.runs[0].font.bold = True
	header_p.runs[0].font.bold = True
	table1.rows[1].cells[0].text = interim
	table1.rows[1].cells[1].text = final
	table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	table1.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

	document.add_page_break()
	overview = document.add_paragraph()
	overview.add_run('Crop(s):  ').font.bold = True
	overview.add_run(', '.join(d['crop']) + '\r\n')
	overview.add_run('Target(s):  ').font.bold = True
	overview.add_run(', '.join(d['target']))
	overview.add_run('\r\nObjective(s): ').font.bold = True
	objectives = document.add_paragraph(d['objective'][0])
	objectives.paragraph_format.left_indent = Inches(0.5)
	
	document.add_page_break()
	header = document.sections[0].header.paragraphs[0]
	run = header.add_run()
	run.add_picture('.\\logo.jpg', width = Inches(1))
	footer = document.sections[0].footer.paragraphs[0]
	footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run2 = footer.add_run()
	if confidential == True:
		run2.text = d['shortname'][0]
		run3 = footer.add_run()
		run3.text = "  (Confidential)"
		run3.font.color.rgb = RGBColor(255,0,0)
	else:
		run2.text = d['shortname'][0]
		run3 = footer.add_run()
		run3.text = "  (Not-Confidential)"
	header.style = document.styles['Header']
	footer.style = document.styles['Footer']
	run.alignment = WD_ALIGN_PARAGRAPH.LEFT
	add_page_number(document.sections[0].header.paragraphs[0])
	document.save(protocol_id + '.docx')


def add_assessments(document, protocol_id, df, font_type, font_size):
	cols = df.columns
	table = document.add_table(rows=1, cols=len(cols))
	table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
	style = document.styles.add_style('table', WD_STYLE_TYPE.PARAGRAPH); font = style.font; font.name = font_type; font.size = Pt(font_size); font.bold = False
	style2 = document.styles.add_style('table_header', WD_STYLE_TYPE.PARAGRAPH); font = style2.font; font.name = font_type; font.size = Pt(font_size); font.bold = True; font.color.rgb = RGBColor(255,255,255)
	hdr_cells = table.rows[0].cells; y = 0
	for x in cols:
		hdr_cells[y].text = x
		p = hdr_cells[y].paragraphs[0]
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.style = document.styles['table_header']
		shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
		hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
		y+=1
	y = 0; index = 1;
	for x in df.iterrows():
		border_row = False
		row = x[1]
		row_cells = table.add_row().cells
		for i in range(len(cols)):
			row_cells[i].text = str(row[i])
			p = row_cells[i].paragraphs[0]
			p.style = document.styles['table']
			set_cell_border(row_cells[i],start={"sz": 3, "val": "single", "color": "#5b6775", "space": "0"})
			set_cell_border(row_cells[i], top={"sz": 3, "val": "single", "color": "#FF0000", "space": "0"})
			if y % 2 == 1:
				shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
				row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
			if index == len(df):
				set_cell_border(row_cells[i], bottom={"sz": 3, "val": "single", "color": "#5b6775", "space": "0"})
		y += 1
		index += 1
	for column in table.columns:
		for cell in column.cells:
		    tc = cell._tc
		    tcPr = tc.get_or_add_tcPr()
		    tcW = tcPr.get_or_add_tcW()
		    tcW.type = 'auto'


def add_trttable(document, protocol_id, df, table_font, font_size):
	cols = df.columns
	table = document.add_table(rows=1,cols=len(cols))
	table.allow_autofit = True
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells; y = 0
	for x in cols:
		hdr_cells[y].text = x
		p = hdr_cells[y].paragraphs[0]
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.style = document.styles['table_header']
		shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
		hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
		y+=1
	y = 0; current = 1; shaded = False
	for x in df.iterrows():
		border_row = False
		row = x[1]
		row_cells = table.add_row().cells
		for i in range(len(cols)):
			row_cells[i].text = str(row[i])
			p = row_cells[i].paragraphs[0]
			p.style = document.styles['table']
			set_cell_border(row_cells[i],start={"sz": 3, "val": "single", "color": "#5b6775", "space": "0"})
			if i == 0:
				if int(row[i]) > current:
					current = int(row[i])
					border_row = True
					if shaded == False:
						shaded = True
					elif shaded == True:
						shaded = False
			if border_row == True:
				set_cell_border(row_cells[i], top={"sz": 3, "val": "single", "color": "#FF0000", "space": "0"})
			if shaded == True:
				shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
				row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
		y += 1
	for column in table.columns:
		for cell in column.cells:
		    tc = cell._tc
		    tcPr = tc.get_or_add_tcPr()
		    tcW = tcPr.get_or_add_tcW()
		    tcW.type = 'auto'
	table.style.name='Table Grid'
	section = document.sections[-1]
	new_width, new_height = section.page_height, section.page_width 
	section.orientation = WD_ORIENT.LANDSCAPE
	section.page_width = new_width
	section.page_height = new_height
	section.left_margin = Inches(0.5)
	section.right_margin = Inches(0.5)
	document.save(protocol_id + '.docx')


def add_instructions(document, protocol_id):
	d = clean_instructions(protocol_id)[2]
	for header in d.keys():
		p = document.add_paragraph()
		p.paragraph_format.keep_with_next = True
		hdr = p.add_run(header)
		content = document.add_paragraph()
		content.paragraph_format.keep_together = True
		for x in d[header]:
			content_run = content.add_run(x)
			content_run.font.size = Pt(12)
			content_run.font.bold = False
		content.paragraph_format.left_indent = Inches(0.5)
		content.paragraph_format.space_after = Pt(6)
		hdr.font.size = Pt(12)
		hdr.font.underline = True
		hdr.font.bold = True
		p.paragraph_format.space_after = Pt(6)
	document.add_page_break()
	document.save(protocol_id + '.docx')


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):

    page_run = paragraph.add_run('\t\t\t\t\t\t\t')
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

"""
treatments_trialdesign_treatedplotaream2 as plotarea, treatments_applicationlist_code as app_code, treatments_trialdesign_statisticaldesign_armdescription as stat_design, treatments_applicationlist_volumemin as app_minvolume, treatments_applicationlist_volumemax as app_maxvolume, treatments_applicationlist_volumeunit_code as app_unit, treatments_applicationlist_requiredmixsize as req_mixsize, treatments_applicationlist_percentageoverage as percentoverage, treatments_applicationlist_mixsizeunit_code as mixsize_unit, treatments_treatmentlist_productamounttotalqty as product_total, treatments_trialdesign_numberofreplicates as replicates 
"""