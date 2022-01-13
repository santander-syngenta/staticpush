from __future__ import print_function
import time
from pprint import pprint
import psycopg2
import numpy as np
from sqlalchemy import create_engine
import pandas as pd
import math
from docx.shared import Cm, Inches, RGBColor, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
import dateutil.parser as parser
from docx.table import _Cell
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pyodbc
from win32com import client
import win32com
import re, datetime
from .models import fpc as FPC, manager, owner, contributor, reviewer, protocol


host = 'deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com'
pwd = 's1030345PASS81997!'
host = 'deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com'
user = 's1030345'
port = '5439'
dbname = 'mio'


host2 = 'usaedmssobdp001.c0or7llb3toc.us-east-1.rds.amazonaws.com'
user2 = 'Prod_read'
dbname2 = 'BioDesign'
port2 = 1433
pwd2 = 'Prod_read@BD'


dbname3 = 'APPLICATION'
host3 = 'usaedwbanp002.syngentaaws.org'
port3 = 1433
user3 = 'READONLY_USER'
pwd3 = 'READONLY_USER'


def simple_connect():
	connection = psycopg2.connect(server=host2,user=user2,password=pwd2,database=dbname2)
	cursor = connection.cursor()
	return connection, cursor


def pd_connect():
	engine = create_engine('postgresql://s1030345:s1030345PASS81997!@deawirbitt001.clwtglrkcnfi.eu-central-1.redshift.amazonaws.com:5439/mio')
	return engine


def bd_pull(protocol_id):
	conn = pyodbc.connect('DRIVER={SQL Server}; SERVER=' + host2 + "; DATABASE=" + dbname2 + '; UID=' + user2 + '; PWD=' + pwd2)
	cursor = conn.cursor()
	p = protocol_id.upper()
	cursor.execute("SELECT ObjectiveFormatted, ConfidentialFormatted, InstructionsFormatted FROM BioDesign.dbo.Protocol WHERE ProtocolIdentifier = '" + p + "';")
	data = cursor.fetchall()[0]
	return data


def ba_pull(shortname):
	conn = pyodbc.connect('DRIVER={SQL Server}; SERVER=' + host3 + "; DATABASE=" + dbname3 + '; UID=' + user3 + '; PWD=' + pwd3)
	cursor = conn.cursor()
	s =  shortname.upper()
	cursor.execute("SELECT top 1 InstructionsRTF from data_access.v_protocol_info where Trial_MasterProtocolID =\'" + s + "\';", )
	data = cursor.fetchall()
	return data


def clean_bd(protocol_id):
	obj = bd_pull(protocol_id)
	objectives = obj[0]
	confidential = obj[1]
	instructions = obj[2]


def data_lake_sync():
	engine = pd_connect()
	print('Connecting to the Data Lake...')
	df = pd.read_sql("SELECT shortname, protocol_id, overview_setup_trialstartyear, status, (overview_users_ownedby_firstname + ' ' + overview_users_ownedby_lastname) as owner, listagg(DISTINCT overview_users_contributors_firstname + ' ' + overview_users_contributors_lastname, ', ') as contributors, listagg(DISTINCT overview_users_reviewers_firstname + ' ' + overview_users_reviewers_lastname, ', ') as reviewers FROM mio.public.mio212_trialprotocol_fact WHERE overview_setup_country_code = 'US' GROUP BY shortname, protocol_id, overview_setup_trialstartyear, status, overview_users_ownedby_firstname, overview_users_ownedby_lastname", engine)
	df2 = pd.read_sql("SELECT derivedpotocolshortname as shortname, listagg(DISTINCT balancing_fpc_firstname + ' ' + balancing_fpc_lastname, ', ') as fpc, listagg(DISTINCT balancing_responsibletrialmanager_firstname + ' ' + balancing_responsibletrialmanager_lastname, ', ') as tr_mgr FROM (SELECT DISTINCT balancing_territory_name, derivedpotocolshortname, balancing_fpc_lastname, balancing_fpc_firstname, balancing_responsibletrialmanager_firstname, balancing_responsibletrialmanager_lastname FROM mio.public.mio189_biobalance_fact) WHERE balancing_territory_name = 'USA' GROUP BY derivedpotocolshortname", engine)
	df = df.dropna(axis = 0, how = 'any', subset = ['shortname','protocol_id'])
	df2 = df2.dropna(axis = 0, how = 'any', subset = ['shortname'])
	output = pd.merge(df, df2, how='left', on=['shortname'])
	output.index = output['shortname']
	output = output.drop(columns=['shortname']).to_dict(orient='index')
	print('Data pulled from Data Lake successfully. Now Comparing to the protocol site data')
	for y in output.keys():
		try:
			protocol.objects.get(shortname = y)
			pass
		except:
			d = check_users(output[y])
			print('Adding protocol: ' , y)
			p = protocol()
			p.shortname = y
			p.protocol_id = d['pid']
			p.year = output[y]['overview_setup_trialstartyear']
			p.status = output[y]['status']
			p.save()
			if d['reviewers'] != []:
				r = reviewer.objects.filter(id__in=d['reviewers'])
				for x in r:
					p.reviewer.add(x)
			if d['contribs'] != [] :
				c = contributor.objects.filter(id__in=d['contribs'])
				for x in c:
					p.contributor.add(x)
			if d['owners'] != []:
				o = owner.objects.filter(id__in=d['owners'])
				for x in o:
					p.owner.add(x)
			if d['fpc'] != []:
				f = FPC.objects.filter(id__in=d['fpc'])
				for x in f:
					p.fpc.add(x)
			if d['mgrs'] != []:
				m = manager.objects.filter(id__in=d['mgrs'])
				for x in m:
					p.tr_mgr.add(x)
			p.save()
	print('Sync Successful')
	return output


def check_users(protocol_object):
	fpcs = protocol_object['fpc']; fpc_indexes = []
	print(protocol_object)
	if fpcs != None and type(fpcs) == str:
		for f in fpcs.split(', '):
			fpc_present = list(FPC.objects.filter(name=f))
			if fpc_present == []:
				print('Adding FPC: ', f)
				f_obj = FPC()
				f_obj.name = f
				f_obj.save()
				fpc_indexes.append(f_obj.id)
			else:
				fpc_indexes.append(fpc_present[0].id)
	tr_mgrs = protocol_object['tr_mgr']; mgr_indexes = []
	if tr_mgrs != None and type(tr_mgrs) == str:
		for t in tr_mgrs.split(', '):
			mgr_present = list(manager.objects.filter(name=t))
			if mgr_present == []:
				print('Adding manager: ', t)
				f_obj = manager()
				f_obj.name = t
				f_obj.save()
				mgr_indexes.append(f_obj.id)
			else:
				mgr_indexes.append(mgr_present[0].id)
	owners = protocol_object['owner']; owner_indexes = []
	if owners != None and type(owners) == str:
		for o in owners.split(', '):
			owner_present = list(owner.objects.filter(name=o))
			if owner_present == []:
				print('Adding Owner: ', o)
				f_obj = owner()
				f_obj.name = o
				f_obj.save()
				owner_indexes.append(f_obj.id)
			else:
				owner_indexes.append(owner_present[0].id)
	contribs = protocol_object['contributors']; contrib_indexes = []
	if contribs != None and type(contribs) == str:
		contribs = contribs.split(', ')
		for c in contribs:
			contrib_present = list(contributor.objects.filter(name=c))
			if contrib_present == []:
				print('Adding Contrib: ', c)
				f_obj = contributor()
				f_obj.name = c
				f_obj.save()
				contrib_indexes.append(f_obj.id)
			else:
				contrib_indexes.append(contrib_present[0].id)
	reviewers = protocol_object['reviewers']; reviewer_indexes = []
	if reviewers != None and type(reviewers) == str:
		for r in reviewers.split(', '):
			reviewer_present = list(reviewer.objects.filter(name=r))
			if reviewer_present == []:
				print('Adding Reviewer: ', r)
				f_obj = reviewer()
				f_obj.name = r
				f_obj.save()
				reviewer_indexes.append(f_obj.id)
			else:
				reviewer_indexes.append(reviewer_present[0].id)
	d = {'fpc':fpc_indexes,'mgrs':mgr_indexes,'owners':owner_indexes,'contribs':contrib_indexes, 'reviewers':reviewer_indexes, 'pid':protocol_object['protocol_id']}
	return d


def pd_overview(protocol_id):
	"""Returns a DataFrame with the relevant fields from the overview table. Runs slower but retains table's column headers."""
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT overview_setup_discipline_name as discipline, overview_setup_project as project, overview_setup_stage_name as stage, overview_setup_country_iso3 as country, overview_setup_origin_code as usage_code, overview_setup_origin_name as usage, overview_users_ownedby_firstname as owner_first, overview_users_ownedby_lastname as owner_last, overview_setup_interimdataneededby as interimdataneededby, overview_setup_finaldataneededby as finaldataneededby, overview_setup_neededby as protocolneededby, overview_objectives_objectivenormal as objective, overview_objectives_objective as overview2, overview_crops_crop_name as crop, overview_users_reviewers_firstname as reviewer_first, overview_users_contributors_firstname as contributor_first, overview_users_contributors_lastname as contributor_last, overview_users_reviewers_lastname as reviewer_last, overview_setup_trialstartyear as start_year, shortname, overview_targets_target_armcode as target_code, overview_targets_target_armdescription as target, overview_crops_crop_code as crop_code FROM mio.public.mio212_trialprotocol_overview WHERE protocol_id =  \'" + p + "\';", engine)
	df2 = pd.read_sql("SELECT balancing_proposedprotocolname FROM mio.public.mio189_biobalance_fact where derivedprotocol_id= \'" + p + "\';", engine)
	df3 = pd.read_sql("SELECT top 1 status, revision, lastupdateddate FROM mio.public.mio212_trialprotocol_fact WHERE protocol_id =  \'" + p + "\';", engine)
	return df, df2, df3


def get_applications(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT treatments_trialdesign_numberofreplicates as no_replicates, treatments_trialdesign_treatedplotaream2 as treated_area,treatments_applicationlist_code as app_code, treatments_applicationlist_volumemin as min_vol, treatments_applicationlist_volumemax as max_vol, treatments_applicationlist_requiredmixsize as mix_size, treatments_applicationlist_percentageoverage as perc_overage, treatments_applicationlist_adjustedmixsize as adjusted_size, treatments_applicationlist_totalarea as total_area, treatments_applicationlist_method_armdescription as app_method, treatments_applicationlist_timing_armdescription as timing, treatments_applicationlist_placement_armdescription as placement, treatments_applicationlist_volumeunit_code as volume_unit, treatments_applicationlist_mixsizeunit_code as mix_unit from mio.public.mio212_trialprotocol_treatments WHERE protocol_id = \'" + p + "\';", engine)
	rows_to_drop = []
	for i, x in df.iterrows():
		if x['app_code'] == None:
			rows_to_drop.append(i)
	df = df.drop(index=rows_to_drop)
	df = df.applymap(str)
	df['mix_size'] = df['mix_size'] + ' ' + df['mix_unit']
	df['min_vol'] = df['min_vol'] + ' ' + df['volume_unit']
	df['max_vol'] = df['max_vol'] + ' ' + df['volume_unit']
	df = df.drop(columns=['mix_unit', 'volume_unit'])
	return df


def get_assessments(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT assessments_standardevaluationname_code as se_code, assessments_standardevaluationname_name as se_name, assessments_partrated_name as part_rated, assessments_ratingtype_code as rating_type, assessments_ratingunit_code as rating_unit, assessments_sampling as sample_size, assessments_samplingunit_code as sample_unit, assessments_collection as collection_basis, assessments_collectionunit_code as collection_unit, assessments_reporting as reporting_basis, assessments_reportingunit_code as reporting_unit, assessments_numberofsubsamples as no_subsamples FROM mio.public.mio212_trialprotocol_assessments where protocol_id =\'" + p + "\';", engine)
	df = remove_blankrows(df)
	df['no_subsamples'] = df['no_subsamples'].fillna(0).map(int)
	df['sample_size'] = df['sample_size'].fillna(0).map(int)
	df['collection_basis'] = df['collection_basis'].fillna(0).map(int)
	df['reporting_basis'] = df['reporting_basis'].fillna(0).map(int)
	df = df.applymap(str)
	df['sample_size'] = df['sample_size'] + '/' + df['sample_unit']
	df['collection_basis'] = df['collection_basis'] + '/' + df['collection_unit']
	df['reporting_basis'] = df['reporting_basis'] + '/' + df['reporting_unit']
	df = df.drop(columns=['sample_unit','collection_unit','reporting_unit'])
	df.columns = df.columns.str.replace('no_subsamples', '# of Sub-\nSamples')
	df.columns = df.columns.str.replace('reporting_basis', 'Reporting\nBasis')
	df.columns = df.columns.str.replace('collection_basis', 'Collection\nBasis')
	df.columns = df.columns.str.replace('sample_size', 'Sample\nSize')
	df.columns = df.columns.str.replace('part_rated', 'Part\nRated')
	df.columns = df.columns.str.replace('se_code', 'SE Code')
	df.columns = df.columns.str.replace('se_name', 'SE Name')
	df.columns = df.columns.str.replace('rating_type', 'Rating\nType')
	df.columns = df.columns.str.replace('rating_unit', 'Rating\nUnit')
	return df


def get_treatments(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT treatments_treatmentlist_treatmentnumber as no, treatments_treatmentlist_treatementseq as seq_no, treatments_treatmentlist_treatementtag as trttag, treatments_treatmentlist_formconcqty as conc_qty, treatments_treatmentlist_rate as rate, treatments_treatmentlist_otherrate as other_rate, treatments_treatmentlist_applcode as appl_code, treatments_treatmentlist_check_name as check_name, treatments_treatmentlist_productamountunit_code as unit, treatments_treatmentlist_treatmenttype_code as treatment_code, treatments_treatmentlist_treatmentname_code as trt_name, treatments_treatmentlist_treatmentname_name as trt_name2, treatments_treatmentlist_formconcunit_code as conc_unit, treatments_treatmentlist_formtype_code as Form_Type, treatments_treatmentlist_rateunit_code as rate_unit, treatments_treatmentlist_otherrateunit_code as otherrate_unit, treatments_treatmentlist_appltiming_armdescription as appl_timing, treatments_treatmentlist_applmethod_armdescription as appl_method, treatments_treatmentlist_applplacement_armdescription as appl_placement, treatments_treatmentlist_treatmenttype_code as trt_type, treatments_treatmentlist_minapplication as min_app, treatments_treatmentlist_productamounttotalqty as total_qty, treatments_treatmentlist_productamountunit_code as total_units, treatments_treatmentlist_treatmentname_name as prod_description, treatments_treatmentlist_applplacement_armcode as placement, treatments_treatmentlist_appltiming_armcode as timing, treatments_treatmentlist_applmethod_armcode as method, treatments_treatmentlist_lotcode as lotcode FROM mio.public.mio212_trialprotocol_treatments WHERE protocol_id = \'" + p + "\';", engine)
	df = remove_blankrows(df)
	return df


def get_instructions(protocol_id):
	engine = pd_connect()
	p = protocol_id.upper()
	df = pd.read_sql("SELECT instructions_instructions as instructions, instructions_confidential as design_codes, instructions_studyruleset_code as strudy_rules FROM mio.public.mio212_trialprotocol_instructions WHERE protocol_id = \'" +p+'\';', engine)
	if df.empty == True:
		return [] 
	elif df.empty == False:
		df = remove_blankrows(df).reset_index()
		df = df.drop(columns=['index'])
		output = {}; c = df.columns
		for x in range(len(df.loc[0])):
			column = []
			for y in range(len(df)):
				if df.iloc[y,x] == None:
					pass
				else:
					column.append(df.iloc[y,x])
			output[c[x]] = column
		return output
	

def clean_instructions(protocol_id):
	d = get_instructions(protocol_id)
	if d == []:
		return [] 
	elif d['instructions'] == []:
		return []
	else:
		instructions = d['instructions'][0]
		occurences = instructions.count('\r\n'); lines = []
		subjects = ['CROPS/SURFACES','CROPS', ' CROPS', 'TARGETS', 'OBJECTIVE','OBJECTIVES','OBJECTIVE(S)','SPECIAL PROTOCOL TASKS','DATA REQUIREMENTS/ESSENTIAL DATA','DATA REQUIEMENTS/ESSENTIAL DATA','EXPERIMENTAL DESIGN AND PLOT DIMENSIONS','EXPERIMENTAL DESIGN & PLOT DIMENSIONS','MAINTANENCE DETAILS', 'MAINTENANCE DETAILS','TREATMENT DETAILS', 'ASSESSMENT TIMING SUMMARY','ASSESSMENT DETAILS','DATA REPORTING DEADLINES','OTHER NOTES','CROP DESTRUCT','DATA REPORTING GUIDELINES','SAFETY AND STEWARDSHIP OF TEST SUBSTANCES IN THIS PROTOCOL','CONTRACT RESEARCH ORGANIZATIONS','CONTRACT RESEARCH ORGANIZATIONS (NOT UNIVERSITIES)', 'CROP DESTRUCT','UNIVERSITIES','DESIGN CODES']
		for x in range(occurences):
			index = instructions.index('\r\n')
			lines.append(instructions[0:index+2])
			instructions = instructions[index+2:]
		y = 0; indexes = []; current_para = []; output = {}; current_hdr = None
		for line in lines:
			try:
				index = line.index(':')
				sub = line[:index]
				header = True
				for x in sub:
					if x.isalpha() == True and x.isupper() == True:
						header = True 
					elif x.isalpha() == True and x.isupper() == False:
						header = False
						break 
					else:
						pass
				if header == True and sub in subjects:
					indexes.append(y)
					output[current_hdr] = current_para
					current_hdr = line
					current_para = []
				else:
					current_para.append(line)
			except:
				if line.strip('\r\n').strip() in subjects:
					indexes.append(y)
					output[current_hdr] = current_para
					current_hdr = line
					current_para = []
				else:
					current_para.append(line)
			y+=1
		output.pop(None)
		out2 = {}
		for x in output.keys():
			lines = output[x]
			l_copy = lines.copy()
			lines_out = []
			for line in l_copy:
				if line == ' \r\n':
					lines.remove(line)
				else:
					if '●' in line:
						line = line.replace('●','•')
					elif '●' in line:
						line = line.replace('●','•')
					lines_out.append(line.replace('\r\n','\n'))
			out2[x] = lines_out
		
		return lines, indexes, out2


def get_overview(protocol_id):
	"""Returns a DataFrame with the relevant fields from the overview table. Runs faster but no column headers."""
	connection, cursor = simple_connect()
	p = protocol_id.upper()
	cursor.execute("SELECT overview_setup_discipline_name as discipline, overview_setup_project as project, overview_setup_stage_name as stage, overview_setup_country_name as country, overview_setup_origin_code as usage_code, overview_setup_origin_name as usage, overview_users_ownedby_firstname as owner_first, overview_users_ownedby_lastname as owner_last, overview_setup_interimdataneededby as interimdataneededby, overview_setup_finaldataneededby as finaldataneededby, overview_setup_neededby as protocolneededby, overview_objectives_objectivenormal as objective, overview_crops_crop_name as crop, overview_users_reviewers_firstname as reviewer_first, overview_users_reviewers_lastname FROM mio.public.mio212_trialprotocol_overview WHERE protocol_id = \'" + p + "\';")
	data = np.array(cursor.fetchall())
	cursor.close(); connection.close()
	return data


def condense_pdoverview(df, df2):
	df['owner'] = df['owner_first'] + ' '+ df['owner_last']
	df['reviewer'] = df['reviewer_first'] + ' '+ df['reviewer_last']
	df['contributor'] = df['contributor_first'] + ' ' + df['contributor_last']
	df = df.drop(columns=['owner_first','owner_last', 'reviewer_first', 'reviewer_last', 'contributor_last', 'contributor_first'])
	output = {}; c = df.columns
	for x in range(len(df.loc[0])):
		column = []
		for y in range(len(df)):
			if (type(df.iloc[y,x]) == float or type(df.iloc[y,x]) == np.float64) and math.isnan(df.iloc[y,x]) == True:
				pass
			elif df.iloc[y,x] == None:
				pass
			else:
				column.append(df.iloc[y,x])
		output[c[x]] = column
	for x in range(len(output['target'])):
		t = output['target'][x] + ' (' + output['target_code'][x] +')'
		output['target'][x] = t
	for x in range(len(output['crop'])):
		c = output['crop'][x] + ' (' + output['crop_code'][x] + ')'
		output['crop'][x] = c
	output.pop('crop_code')
	output.pop('target_code')
	if df2.empty == False:
		output['title'] = df2.iloc[0,0]
	else:
		output['title'] = ' '
	return output


def remove_blankrows(df):
	rows = []
	for x in range(len(df)):
		blank = True
		for y in range(len(df.iloc[x])):
			if (type(df.iloc[x,y]) == float and math.isnan(df.iloc[x,y]) == True) or (type(df.iloc[x,y]) == np.float64 and math.isnan(df.iloc[x,y]) == True) or df.iloc[x,y] == None:
				pass
			else:
				blank = False
		if blank == True:
			rows.append(x)
	df = df.drop(index=rows)
	return df


def clean_trtdf(protocol_id):
	df = get_treatments(protocol_id)
	df = remove_blankrows(df)
	df.columns = df.columns.str.replace('trt_name','Treatment Name')
	df2 = df[['Treatment Name','total_qty','total_units','Treatment Name2']]
	df = df.fillna(' ')
	df = df.applymap(str)
	for x in range(len(df)):
		df.iloc[x,0] = str(int(float(df.iloc[x,0])))

	"""Rename & Reorder columns in Treatment Table"""
	df.columns = df.columns.str.replace('no','No.')
	df = df.drop(columns=['trt_type'])
	df.columns = df.columns.str.replace('appl_code','Code')
	df.columns = df.columns.str.replace('form_type', 'Type')
	df.columns = df.columns.str.replace('min_app', 'Min #\nAppl')
	df.columns = df.columns.str.replace('other_rate','Other\nRate')
	df.columns = df.columns.str.replace('otherrate_unit', 'Other Rate\nUnit')
	df.columns = df.columns.str.replace('rate', 'Rate')
	df.columns = df.columns.str.replace('Rate_unit','Rate\nUnit')
	df.columns = df.columns.str.replace('conc_unit', 'Form.\nUnit')
	df.columns = df.columns.str.replace('conc_qty', 'Form.')
	df.columns = df.columns.str.replace('lotcode', 'Lot\nCode')
	df.columns = df.columns.str.replace('method','Method')
	df.columns = df.columns.str.replace('placement', 'Place\nment')
	df.columns = df.columns.str.replace('timing', 'Timing')
	df.columns = df.columns.str.replace('prod_description', 'Description')
	df = df[['No.', 'trttag', 'seq_No.', 'Treatment Name', 'Description','Form.', 'Form.\nUnit', 'Type', 'Rate','Rate\nUnit','Other\nRate','Other Rate\nUnit','Place\nment','Timing','Method','Min #\nAppl','Code','Lot\nCode']]
	remove = []
	for column in df.columns:
		blank = True
		for value in df[column]:
			if value.strip():
				blank = False
				break
		if blank == True:
			remove.append(column)
	df = df.drop(columns=remove)
	df2 = df2.dropna(subset=['total_qty'])
	unique = df2['Treatment Name'].unique(); d = {}
	for x in unique:
		d[x] = [0,None,None]
	for x in df2.iterrows():
		d[x[1][0]][0] += round(float(x[1][1]),2)
		d[x[1][0]][1] = x[1][2]
		d[x[1][0]][2] = x[1][3]
	df['No.'] = df['No.'].astype(int)
	df = df.sort_values(by=['No.', 'trttag', 'seq_No.'])
	df = df.drop(columns=['seq_No.','trttag'])
	df = df.dropna(axis=1,how='all')
	"""Removing Trailing zeros without rounding"""
	df['Rate'] = df['Rate'].str.replace('.0$','')
	df['Form.'] = df['Form.'].str.replace('.0$','')
	df['Min #\nAppl'] = df['Min #\nAppl'].str.replace('.0$','')

	return df, d


def printer(protocol_id):
	print_doc(protocol_id, 'Arial', 12, True, False)


def print_doc(protocol_id, shortname, font, font_size, confidential, color):
	df, totals = clean_trtdf(protocol_id)
	df2, df3, status_df = pd_overview(protocol_id)
	if df2.empty == False:
		overview = condense_pdoverview(df2,df3)
	else:
		overview = []
	assessments = get_assessments(protocol_id)
	instructions = clean_instructions(protocol_id)
	d = Document()
	if overview != []:
		add_cover(d, protocol_id, overview, confidential, font, font_size, status_df)
	if instructions != []:
		instructions = instructions[2]
		add_instructions2(d, instructions, protocol_id,font_size,font)
	d.add_section(WD_SECTION.NEW_PAGE)
	if assessments.empty == False:
		add_assessments(d, protocol_id, assessments, font, font_size, color)
	if totals != {}:
		add_totals(d, totals, font, font_size, confidential, color)
	if df.empty == False: 
		d.add_page_break()
		add_trttable(d, protocol_id, df, font, font_size, color, confidential)
	
	d.save('C:\\inetpub\\wwwroot\\nematool\\static\\docs\\' + shortname + '.docx')
	


def print_doc2(protocol_id, font, font_size, confidential, color):
	###Data pull from Data Lake
	df, totals = clean_trtdf(protocol_id)
	df2, df3 = pd_overview(protocol_id)
	instructions = bd_pull(protocol_id)
	if df2.empty == False:
		overview = condense_pdoverview(df2,df3)
	else:
		overview = []
	assessments = get_assessments(protocol_id)
	
	d = Document()

	if overview != []:
		add_cover(d, protocol_id, overview, confidential, font, font_size)

	##Add Instructions
	parser = HtmlToDocx()
	obj = bd_pull(protocol_id)
	objectives = obj[0]
	confidential = obj[1]
	instructions = obj[2]
	parser.add_html_to_document(instructions, d)

	##Add tables sections
	d.add_section(WD_SECTION.NEW_PAGE)
	if assessments.empty == False:
		add_assessments(d, protocol_id, assessments, font, font_size, color)
		d.add_page_break()
	if totals != {}:
		add_totals(d, totals, font, font_size, confidential, color)
		d.add_page_break()
	if df.empty == False: 
		add_trttable(d, protocol_id, df, font, font_size, color)
	if overview != []:
		d.save(overview['shortname'][0] + '.docx')
	else:
		d.save(protocol_id + '.docx')


def inst(instructions):
	sub = instructions; divs = []
	for i in range(instructions.count('<div')):
		start = sub.index('<div')
		end  = sub.index('</div>')
		div = sub[start : end + 6]
		divs.append(div)
		sub = sub[end+6:]
	return divs

def add_cover(document, protocol_id, d, confidential, font_type, font_size, status_df):
	document.styles['Normal'].font.name = font_type
	shortname = document.add_paragraph()
	short_run = shortname.add_run("\r\n\r\n" + d['shortname'][0])
	short_run.font.bold = True
	short_run.font.size = Pt(font_size + 2)
	shortname.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	interim, final  = parser.parse(d['interimdataneededby'][0]).strftime("%a, %B %d %Y"), parser.parse(d['finaldataneededby'][0]).strftime("%a, %B %d %Y")
	cover_body = document.add_paragraph()
	title = cover_body.add_run(d['title'])
	title.font.size = Pt(font_size)
	title.font.bold = True

	info = cover_body.add_run('\r\nProtocol Owner: ' + d['owner'][0] + '\r\nContributors: ' + ', '.join(d['contributor']) + '\r\nReviewers: ' + ', '.join(d['reviewer']))
	cover_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	table1 = document.add_table(rows=3,cols=2)
	table1.style.name='Table Grid'
	table1.alignment = WD_TABLE_ALIGNMENT.CENTER
	header_row = table1.rows[0].cells
	header_row[0].text = "Interim Data Needed By:"
	header_row[1].text = "Final Data Needed By:"
	header_p = header_row[0].paragraphs[0]
	header_p2 = header_row[1].paragraphs[0]
	header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
	header_p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
	header_p2.runs[0].font.bold = True
	header_p.runs[0].font.bold = True
	table1.rows[1].cells[0].text = interim
	table1.rows[1].cells[1].text = final
	table1.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	table1.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	try:
		status = status_df.iloc[0,0]
	except:
		status = 'N/A'
	try:
		revision = status_df.iloc[0,1]
	except:
		revision = 'N/A'
	try:
		last_update = status_df.iloc[0,2].replace(re.findall(r'.\d{7}$', status_df.iloc[0,2])[0], '')
		last_update_date = datetime.datetime.strptime(last_update, "%Y-%m-%dT%H:%M:%S").strftime("%b %w, %Y; %I:%M%p")
	except:
		last_update_date = 'N/A'
	table2 = document.add_table(rows=2,cols=3)
	table2.style.name='Table Grid'
	table2.alignment = WD_TABLE_ALIGNMENT.CENTER
	header_row = table2.rows[0].cells
	header_row[0].text = "Protocol Status:"
	header_row[1].text = "Revision No.:"
	header_row[2].text = "Last Updated:"
	header_p = header_row[0].paragraphs[0]
	header_p2 = header_row[1].paragraphs[0]
	header_p3 = header_row[2].paragraphs[0]
	header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	header_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	header_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
	header_p2.runs[0].font.bold = True
	header_p3.runs[0].font.bold = True
	header_p.runs[0].font.bold = True
	table2.rows[1].cells[0].text = status
	table2.rows[1].cells[1].text = revision
	table2.rows[1].cells[2].text = last_update_date
	table2.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	table2.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	table2.rows[1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
	"""
	document.add_page_break()
	overview = document.add_paragraph()
	overview.add_run('Crop(s):  ').font.bold = True
	overview.add_run(', '.join(d['crop']) + '\r\n')
	overview.add_run('Target(s):  ').font.bold = True
	overview.add_run(', '.join(d['target']))
	overview.add_run('\r\nObjective(s): ').font.bold = True
	objectives = document.add_paragraph(d['objective'][0])
	objectives.paragraph_format.left_indent = Inches(0.5)
	"""
	document.add_page_break()
	header = document.sections[0].header.paragraphs[0]
	run = header.add_run()
	run.add_picture('.\\protocols\\logo.jpg', width = Inches(1))
	footer = document.sections[0].footer.paragraphs[0]
	footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	run2 = footer.add_run()
	if confidential == True:
		run2.text = d['shortname'][0]
		run3 = footer.add_run()
		run3.text = "  For Internal Distribution (Confidential)"
		run3.font.color.rgb = RGBColor(255,0,0)
	else:
		run2.text = d['shortname'][0]
		run3 = footer.add_run()
		run3.text = "  Suitable for External Distribution (Confidential)"
		run3.font.color.rgb = RGBColor(255,0,0)
	header.style = document.styles['Header']
	footer.style = document.styles['Footer']
	run.alignment = WD_ALIGN_PARAGRAPH.LEFT
	add_page_number(document.sections[0].footer.paragraphs[0])


def add_assessments(document, protocol_id, df, font_type, font_size, color):
	title = document.add_paragraph()
	title_run = title.add_run('Assessments')
	title_run.font.bold = True
	title_run.font.size = Pt(font_size + 2)
	title_run.font.underline = True
	title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	cols = df.columns
	table = document.add_table(rows=1, cols=len(cols))
	table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
	if font_type == 'Arial':
		font_type2 = 'Arial Narrow'
	else:
		font_type2 = font_type
	style = document.styles.add_style('table', WD_STYLE_TYPE.PARAGRAPH); font = style.font; font.name = font_type2; font.size = Pt(font_size); font.bold = False
	style2 = document.styles.add_style('table_header', WD_STYLE_TYPE.PARAGRAPH); font = style2.font; font.name = font_type; font.size = Pt(font_size); font.bold = True; font.color.rgb = RGBColor(255,255,255)
	hdr_cells = table.rows[0].cells; y = 0
	for x in cols:
		hdr_cells[y].text = x
		p = hdr_cells[y].paragraphs[0]
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.style = document.styles['table_header']
		shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
		hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
		y+=1
	y = 0; index = 1;
	for x in df.iterrows():
		border_row = False
		row = x[1]
		row_cells = table.add_row().cells
		for i in range(len(cols)):
			row_cells[i].text = str(row[i])
			p = row_cells[i].paragraphs[0]
			p.style = document.styles['table']
			set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			if y % 2 == 1:
				if color == True:
					shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
				else:
					shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
				row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
			if index == len(df):
				set_cell_border(row_cells[i], bottom={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			if i == len(cols)-1:
				set_cell_border(row_cells[i], end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
		y += 1
		index += 1
	for column in table.columns:
		for cell in column.cells:
		    tc = cell._tc
		    tcPr = tc.get_or_add_tcPr()
		    tcW = tcPr.get_or_add_tcW()
		    tcW.type = 'auto'


def add_trttable(document, protocol_id, df, table_font, font_size, color, confidential):
	if confidential == False:
		df = df.drop(columns = ['Description'])
	title = document.add_paragraph()
	title_run = title.add_run('Treatments')
	title_run.font.bold = True
	title_run.font.size = Pt(font_size + 2)
	title_run.font.underline = True
	title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	cols = df.columns
	table = document.add_table(rows=1,cols=len(cols))
	table.allow_autofit = True
	table.alignment = WD_TABLE_ALIGNMENT.CENTER
	hdr_cells = table.rows[0].cells; y = 0
	for x in cols:
		hdr_cells[y].text = x
		p = hdr_cells[y].paragraphs[0]
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		p.style = document.styles['table_header']
		shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
		hdr_cells[y]._tc.get_or_add_tcPr().append(shading_elm_1)
		y+=1
	y = 0; current = 1; shaded = False
	for x in df.iterrows():
		border_row = False
		row = x[1]
		row_cells = table.add_row().cells
		last = False
		if y == len(df)-1:
			last = True
		for i in range(len(cols)): 
			row_cells[i].text = str(row[i])
			p = row_cells[i].paragraphs[0]
			p.style = document.styles['table']
			set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},)
			if i == 0:
				if int(row[i]) > current:
					current = int(row[i])
					border_row = True
					if shaded == False:
						shaded = True
					elif shaded == True:
						shaded = False
			elif i == len(cols)-1:
				set_cell_border(row_cells[i],end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			if border_row == True:
				set_cell_border(row_cells[i], top={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
			if shaded == True:
				if color == True:
					shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
				else:
					shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
				row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
			if last == True:
				set_cell_border(row_cells[i],bottom={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
		y += 1
	for column in table.columns:
		for cell in column.cells:
		    tc = cell._tc
		    tcPr = tc.get_or_add_tcPr()
		    tcW = tcPr.get_or_add_tcW()
		    tcW.type = 'auto'
	table.style.name='Table Grid'
	section = document.sections[-1]
	new_width, new_height = section.page_height, section.page_width 
	section.orientation = WD_ORIENT.LANDSCAPE
	section.page_width = new_width
	section.page_height = new_height
	section.left_margin = Inches(0.5)
	section.right_margin = Inches(0.5)
	section.top_margin = Inches(0.5)
	section.bottom_margin = Inches(0.5)
	section.header_distance = Inches(0.35)
	section.footer_distance = Inches(0.35)


def add_totals(document, d, font, font_size, confidential, color):
	"""Removes confidential trt_name info and makes table"""
	title = document.add_paragraph()
	title_run = title.add_run('Product Totals')
	title_run.font.bold = True
	title_run.font.size = Pt(font_size + 2)
	title_run.font.underline = True
	title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	if confidential == False:
		data = {}
		for x in d.keys():
			data[x] = d[x][0:2]
		table = document.add_table(rows=1,cols=3)
		table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
		hdr_cells = table.rows[0].cells; headers = ['Treatment Name','Product\nTotal','Unit']
		for x in range(3):
			hdr_cells[x].text = headers[x]
			p = hdr_cells[x].paragraphs[0]
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p.style = document.styles['table_header']
			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
			hdr_cells[x]._tc.get_or_add_tcPr().append(shading_elm_1)
			set_cell_border(hdr_cells[x], end={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"},start={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"})
		y = 0
		for x in data.keys():
			row_cells = table.add_row().cells
			row_cells[0].text = x
			row_cells[1].text = str(round(data[x][0],2))
			row_cells[2].text = data[x][1]
			for i in range(3):
				p = row_cells[i].paragraphs[0]
				p.style = document.styles['table']
				set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if y % 2 == 1:
					if color == True:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
					else:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
					row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
			y+=1
		for column in table.columns:
			for cell in column.cells:
			    tc = cell._tc
			    tcPr = tc.get_or_add_tcPr()
			    tcW = tcPr.get_or_add_tcW()
			    tcW.type = 'auto'
	else:
		data ={}
		for x in d.keys():
			data[x] = d[x]
		table = document.add_table(rows=1,cols=4)
		table.allow_autofit = True; table.alignment = WD_TABLE_ALIGNMENT.CENTER;
		hdr_cells = table.rows[0].cells; headers = ['Treatment Name','Alternate','Product\nTotal','Unit']
		for x in range(4):
			hdr_cells[x].text = headers[x]
			p = hdr_cells[x].paragraphs[0]
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			p.style = document.styles['table_header']
			shading_elm_1 = parse_xml(r'<w:shd {} w:fill="5b6775"/>'.format(nsdecls('w')))
			hdr_cells[x]._tc.get_or_add_tcPr().append(shading_elm_1)
			set_cell_border(hdr_cells[x], end={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"},start={"sz": 1, "val": "single", "color": "#ffffff", "space": "0"})
		y = 0
		for x in data.keys():
			row_cells = table.add_row().cells
			row_cells[0].text = x
			row_cells[1].text = data[x][2]
			row_cells[2].text = str(round(data[x][0],2))
			row_cells[3].text = data[x][1]
			for i in range(4):
				p = row_cells[i].paragraphs[0]
				p.style = document.styles['table']
				set_cell_border(row_cells[i],start={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"},end={"sz": 1, "val": "single", "color": "#5b6775", "space": "0"})
				if y % 2 == 1:
					if color == True:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="b5d4ff"/>'.format(nsdecls('w')))
					else:
						shading_elm_1 = parse_xml(r'<w:shd {} w:fill="c9c9c9"/>'.format(nsdecls('w')))
					row_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
			y+=1
		for column in table.columns:
			for cell in column.cells:
			    tc = cell._tc
			    tcPr = tc.get_or_add_tcPr()
			    tcW = tcPr.get_or_add_tcW()
			    tcW.type = 'auto'


def add_instructions(document, instructions, protocol_id,font_size, font):
	d = instructions
	for header in d.keys():
		p = document.add_paragraph()
		p.paragraph_format.keep_with_next = True
		try:
			index = header.index(':')
			f = header[:index]
			s = header[index:]
			hdr = p.add_run(f)
			hdr2 = p.add_run(s.strip('\r\n'))
		except:
			hdr = p.add_run(header.strip('\r\n') + ':')
		content = document.add_paragraph()
		content.paragraph_format.keep_together = True; y = 1;
		for x in d[header]:
			if y != len(d[header]):
				try:
					index = x[:50].index(':')
					bold = x[:index]
					sub = x[index:]
					bold_run = content.add_run(bold)
					sub_run = content.add_run(sub)
					bold_run.font.bold = True
					sub_run.font.bold = False
					bold_run.font.size = Pt(font_size)
					sub_run.font.size = Pt(font_size)
				except:
					content_run = content.add_run(x)
					content_run.font.size = Pt(font_size)
					content_run.font.bold = False
			else:
				try:
					index = x[:50].index(':')
					bold = x[:index]
					sub = x[index:]
					bold_run = content.add_run(bold)
					sub_run = content.add_run(sub.strip('\r\n'))
					bold_run.font.bold = True
					sub_run.font.bold = False
					bold_run.font.size = Pt(font_size)
					sub_run.font.size = Pt(font_size)
				except:
					content_run = content.add_run(x.strip('\r\n'))
					content_run.font.size = Pt(font_size)
					content_run.font.bold = False
			y += 1
		content.paragraph_format.left_indent = Inches(0.15)
		content.paragraph_format.space_after = Pt(10)
		content.paragraph_format.line_spacing = 1
		hdr.font.size = Pt(font_size + 2)
		hdr.font.underline = True
		hdr.font.bold = True
		p.paragraph_format.space_after = Pt(6)
	section = document.sections[-1]
	section.left_margin = Inches(0.35)
	section.right_margin = Inches(0.35)
	section.top_margin = Inches(0.5)
	section.bottom_margin = Inches(0.5)
	section.header_distance = Inches(0.35)
	section.footer_distance = Inches(0.35)
	document.add_page_break()


def add_instructions2(document, instructions, protocol_id,font_size, font):
	d = instructions
	for header in d.keys():
		content = document.add_paragraph()
		try:
			index = header.index(':')
			f = header[:index]
			s = header[index:]
			hdr = content.add_run(f)
			hdr2 = content.add_run(s.replace('\r\n','\n'))
		except:
			hdr = content.add_run(header.strip('\r\n') + ':\n')
		content.paragraph_format.keep_together = True; y = 1;
		for x in d[header]:
			if y != len(d[header]):
				try:
					index = x[:50].index(':')
					bold = x[:index]
					sub = x[index:]
					bold_run = content.add_run(bold)
					sub_run = content.add_run(sub)
					bold_run.font.bold = True
					sub_run.font.bold = False
					bold_run.font.size = Pt(font_size)
					sub_run.font.size = Pt(font_size)
				except:
					content_run = content.add_run(x)
					content_run.font.size = Pt(font_size)
					content_run.font.bold = False
			else:
				try:
					index = x[:50].index(':')
					bold = x[:index]
					sub = x[index:]
					bold_run = content.add_run(bold)
					sub_run = content.add_run(sub.strip('\r\n'))
					bold_run.font.bold = True
					sub_run.font.bold = False
					bold_run.font.size = Pt(font_size)
					sub_run.font.size = Pt(font_size)
				except:
					content_run = content.add_run(x.strip('\r\n'))
					content_run.font.size = Pt(font_size)
					content_run.font.bold = False
			y += 1
		content.paragraph_format.left_indent = Inches(0)
		content.paragraph_format.space_after = Pt(6)
		content.paragraph_format.line_spacing = 1
		hdr.font.size = Pt(font_size + 2)
		hdr.font.underline = True
		hdr.font.bold = True
	section = document.sections[-1]
	section.left_margin = Inches(0.35)
	section.right_margin = Inches(0.35)
	section.top_margin = Inches(0.5)
	section.bottom_margin = Inches(0.5)
	section.header_distance = Inches(0.35)
	section.footer_distance = Inches(0.35)
	document.add_page_break()


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
 
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
 
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
 
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
 
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):

    page_run = paragraph.add_run('    ')
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)


def ConvertRtfToDocx(rootDir, file):
    word = win32com.client.Dispatch("Word.Application")
    wdFormatDocumentDefault = 16
    wdHeaderFooterPrimary = 1
    doc = word.Documents.Open(rootDir + "\\" + file)
    for pic in doc.InlineShapes:
        pic.LinkFormat.SavePictureWithDocument = True
    for hPic in doc.sections(1).headers(wdHeaderFooterPrimary).Range.InlineShapes:
        hPic.LinkFormat.SavePictureWithDocument = True
    doc.SaveAs(str(rootDir + "\\refman.docx"), FileFormat=wdFormatDocumentDefault)
    doc.Close()
    word.Quit()
	
"""
treatments_trialdesign_treatedplotaream2 as plotarea, treatments_applicationlist_code as app_code, treatments_trialdesign_statisticaldesign_armdescription as stat_design, treatments_applicationlist_volumemin as app_minvolume, treatments_applicationlist_volumemax as app_maxvolume, treatments_applicationlist_volumeunit_code as app_unit, treatments_applicationlist_requiredmixsize as req_mixsize, treatments_applicationlist_percentageoverage as percentoverage, treatments_applicationlist_mixsizeunit_code as mixsize_unit, treatments_treatmentlist_productamounttotalqty as product_total, treatments_trialdesign_numberofreplicates as replicates 
"""
